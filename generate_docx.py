#!/usr/bin/env python3
"""Generate DOCX files for V6 Blog Post and Implementation Guide."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


def set_style(doc):
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_code_block(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = val
    return table


def generate_blog_post():
    doc = Document()
    set_style(doc)

    doc.add_heading('Automating AWS Backup failure investigations\nwith Amazon DevOps Agent', 0)
    doc.add_paragraph('Anil Kukkunur, Solutions Architect, AWS').italic = True
    doc.add_paragraph()

    doc.add_paragraph(
        'When managing backups across dozens or hundreds of accounts in AWS Organizations, '
        'investigating failed backup jobs can be time-consuming. Each failure might involve '
        'different root causes—from AWS KMS key access issues to misconfigured IAM policies—'
        'and tracing the problem across accounts requires context-switching and manual log analysis.'
    )
    doc.add_paragraph(
        'In this post, I show you how to build an automated investigation pipeline that detects '
        'failed backup and copy jobs across your entire organization and uses Amazon DevOps Agent '
        'to autonomously perform root cause analysis, delivering findings directly to your team\'s '
        'Slack channel.'
    )

    doc.add_paragraph('Content level: 300 – Advanced')
    doc.add_paragraph('Time to deploy: ~10 minutes')

    # Overview
    add_heading(doc, 'Overview')
    doc.add_paragraph(
        'AWS Backup is a fully managed backup service that centralizes and automates data protection '
        'across AWS services. When you designate a delegated administrator account, you gain '
        'cross-account visibility into backup jobs across your organization.'
    )
    doc.add_paragraph(
        'This solution takes that visibility a step further: when any backup or copy job fails in '
        'any member account, an automated investigation is triggered and root cause analysis is '
        'delivered to Slack within minutes—without human intervention.'
    )

    # Architecture
    add_heading(doc, 'Architecture')
    doc.add_paragraph(
        'The solution uses a two-account deployment model: a one-time setup in the management account, '
        'and a single deployment script in the delegated admin that handles everything else—including '
        'automatic rollout to all member accounts via AWS CloudFormation StackSets.'
    )
    doc.add_paragraph('[Architecture diagram placeholder]').italic = True
    doc.add_paragraph('The flow works as follows:')
    items = [
        'A backup or copy job fails in a member account.',
        'An Amazon EventBridge rule in the member account (deployed via StackSet) captures the failure event and forwards it to the delegated admin account\'s event bus.',
        'An EventBridge rule in the delegated admin invokes an AWS Lambda function.',
        'The Lambda function transforms the event into a signed webhook payload and sends it to the DevOps Agent.',
        'DevOps Agent assumes a cross-account investigation role into the member account, analyzes the failure, and delivers root cause analysis to your Slack channel.',
    ]
    for i, item in enumerate(items, 1):
        doc.add_paragraph(f'{i}. {item}')

    doc.add_paragraph('Key design decisions:', style='List Bullet')
    for d in [
        'Event-driven — zero cost when nothing fails, instant detection when something does.',
        'Fully automated — the deploy script creates the Agent Space, webhook, main stack, and StackSet in one run.',
        'No member account logins — SERVICE_MANAGED StackSets deploy via the Organizations trust.',
        'Auto-scaling coverage — new accounts joining the OU are automatically covered.',
    ]:
        doc.add_paragraph(d, style='List Bullet')

    # Prerequisites
    add_heading(doc, 'Prerequisites')
    add_table(doc,
        ['Requirement', 'Details'],
        [
            ['AWS Organizations', 'With delegated administrator for AWS Backup'],
            ['Management account access', 'One-time setup only'],
            ['Amazon DevOps Agent', 'GA access'],
            ['AWS CLI v2', 'Installed and configured'],
            ['Slack workspace', 'Permission to install apps'],
        ]
    )

    # Step 1
    add_heading(doc, 'Step 1: One-time setup (management account)')
    doc.add_paragraph(
        'Register your delegated admin account for CloudFormation StackSets. This allows the '
        'delegated admin to deploy resources to member accounts without needing management account '
        'credentials again.'
    )
    doc.add_paragraph('Run the following from the management account:')
    add_code_block(doc,
        'aws organizations register-delegated-administrator \\\n'
        '  --account-id YOUR_DELEGATED_ADMIN_ACCOUNT_ID \\\n'
        '  --service-principal member.org.stacksets.cloudformation.amazonaws.com'
    )
    doc.add_paragraph('Verify:')
    add_code_block(doc,
        'aws organizations list-delegated-administrators \\\n'
        '  --service-principal member.org.stacksets.cloudformation.amazonaws.com'
    )

    # Step 2
    add_heading(doc, 'Step 2: Deploy the solution (delegated admin account)')
    doc.add_paragraph('Clone the repository and run the deployment script:')
    add_code_block(doc,
        'git clone https://github.com/kukkunuruanil/aws-backup-devops-agent-v6.git\n'
        'cd aws-backup-devops-agent\n'
        './deploy.sh'
    )
    doc.add_paragraph('The script prompts for 3 inputs:')
    add_table(doc,
        ['Prompt', 'Example'],
        [
            ['Organization ID', 'o-abc1234567'],
            ['Target OU ID', 'ou-xxxx-xxxxxxxx or r-xxxx (root)'],
            ['Region', 'us-west-2'],
        ]
    )
    doc.add_paragraph()
    doc.add_paragraph('The script then performs five automated steps:')
    add_table(doc,
        ['Step', 'Action'],
        [
            ['1/5', 'Creates DevOps Agent Space & HMAC Webhook'],
            ['2/5', 'Deploys main CloudFormation stack (Lambda, EventBridge, IAM, Secrets)'],
            ['3/5', 'Associates AWS account with Agent Space'],
            ['4/5', 'Deploys StackSet to all member accounts (event forwarder + investigation role)'],
            ['5/5', 'Verifies deployment with per-account status table + Lambda test'],
        ]
    )

    doc.add_paragraph()
    doc.add_paragraph('Expected verification output:')
    add_code_block(doc,
        '┌─────────────────┬──────────┬───────────────────────┐\n'
        '│ Account         │ Region   │ Status                │\n'
        '├─────────────────┼──────────┼───────────────────────┤\n'
        '│ 444455556666    │ us-west-2│ ✓ SUCCEEDED           │\n'
        '│ 777788889999    │ us-west-2│ ✓ SUCCEEDED           │\n'
        '└─────────────────┴──────────┴───────────────────────┘\n'
        '\n'
        '→ Testing Lambda (sending synthetic failure event)...\n'
        '✓ Lambda invocation: HTTP 200 | Response: {"statusCode": 200}'
    )

    # Step 3
    add_heading(doc, 'Step 3: Connect Slack')
    doc.add_paragraph('1. In the DevOps Agent console, navigate to Capability Providers and register Slack.')
    doc.add_paragraph('2. Complete the OAuth authorization in your Slack workspace.')
    doc.add_paragraph('3. Associate the Slack channel with your agent space:')
    add_code_block(doc,
        'aws devops-agent associate-service \\\n'
        '  --agent-space-id YOUR_AGENT_SPACE_ID \\\n'
        '  --service-id YOUR_SLACK_SERVICE_UUID \\\n'
        '  --configuration \'{"slack":{"workspaceId":"YOUR_TEAM_ID",...}}\' \\\n'
        '  --region us-west-2'
    )
    doc.add_paragraph('4. Add the DevOps Agent bot to your channel: /invite @AWS DevOps Agent')

    # What gets deployed
    add_heading(doc, 'What gets deployed')
    add_table(doc,
        ['Account', 'Resource', 'Purpose'],
        [
            ['Delegated admin', 'Agent Space', 'Receives webhook events'],
            ['Delegated admin', 'HMAC Webhook', 'Authenticated event intake'],
            ['Delegated admin', 'Lambda BackupFailureBridge', 'Transforms events → webhook'],
            ['Delegated admin', 'EventBridge rule', 'Catches FAILED events'],
            ['Delegated admin', 'EventBus policy', 'Accepts org events'],
            ['Delegated admin', 'IAM roles (2)', 'Agent + Lambda execution'],
            ['Delegated admin', 'Secrets Manager secret', 'Webhook credentials'],
            ['Member accounts', 'EventBridge rule', 'Forwards FAILED events'],
            ['Member accounts', 'IAM role (forwarding)', 'Cross-account event delivery'],
            ['Member accounts', 'IAM role (investigation)', 'Agent cross-account analysis'],
        ]
    )

    # Cleanup
    add_heading(doc, 'Cleanup')
    add_code_block(doc,
        '# Remove StackSet instances\n'
        'aws cloudformation delete-stack-instances \\\n'
        '  --stack-set-name BackupEventForwarder \\\n'
        '  --deployment-targets OrganizationalUnitIds=YOUR_OU_ID \\\n'
        '  --regions us-west-2 --no-retain-stacks \\\n'
        '  --call-as DELEGATED_ADMIN --region us-west-2\n\n'
        '# Delete StackSet\n'
        'aws cloudformation delete-stack-set \\\n'
        '  --stack-set-name BackupEventForwarder \\\n'
        '  --call-as DELEGATED_ADMIN --region us-west-2\n\n'
        '# Delete main stack\n'
        'aws cloudformation delete-stack --stack-name BackupDevOpsAgent --region us-west-2'
    )

    # Conclusion
    add_heading(doc, 'Conclusion')
    doc.add_paragraph(
        'In this post, you built an automated backup failure investigation pipeline that covers '
        'your entire AWS Organization in two steps: register the delegated admin for StackSets '
        '(one-time), then run a single script that creates everything with full verification. '
        'The solution costs nothing when backups succeed and automatically covers new accounts.'
    )

    doc.add_paragraph()
    p = doc.add_paragraph('About the author')
    p.runs[0].bold = True
    doc.add_paragraph(
        'Anil Kukkunur is a Solutions Architect at AWS, focused on helping customers build '
        'resilient data protection strategies using AWS Backup and AWS Organizations.'
    )

    path = os.path.join(OUTPUT_DIR, 'Blog_Post_AWS_Backup_DevOps_Agent.docx')
    doc.save(path)
    print(f'✓ Blog post: {path}')
    return path


def generate_implementation_guide():
    doc = Document()
    set_style(doc)

    doc.add_heading('Implementation Guide\nAWS Backup DevOps Agent', 0)
    doc.add_paragraph('Automated Deployment — Delegated Admin Account').italic = True
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('Time Required: ').bold = True
    p.add_run('~10 minutes')
    p = doc.add_paragraph()
    p.add_run('Deploy from: ').bold = True
    p.add_run('Delegated admin account only')
    p = doc.add_paragraph()
    p.add_run('Member account work: ').bold = True
    p.add_run('None (StackSet handles it automatically)')

    # Architecture
    add_heading(doc, 'Architecture')
    add_code_block(doc,
        '┌─────────────────────────────────────────────────────────────────────┐\n'
        '│                        AWS Organization                             │\n'
        '│                                                                     │\n'
        '│  ┌────────────────────┐      ┌──────────────────────────────────┐  │\n'
        '│  │ Member Account     │      │ Delegated Admin Account          │  │\n'
        '│  │                    │      │                                  │  │\n'
        '│  │ Backup FAILS       │      │ EventBridge → Lambda → Webhook   │  │\n'
        '│  │      │             │      │      ↑                    │      │  │\n'
        '│  │      ▼             │      │      │                    ▼      │  │\n'
        '│  │ EventBridge ───────┼──────┼──────┘            DevOps Agent   │  │\n'
        '│  │ (auto-deployed)    │      │                    │    │        │  │\n'
        '│  │                    │      │                    │    ▼        │  │\n'
        '│  │ Investigation ◄────┼──────┼────────────────────┘  Slack     │  │\n'
        '│  │ Role (auto)        │      │                                  │  │\n'
        '│  └────────────────────┘      └──────────────────────────────────┘  │\n'
        '└─────────────────────────────────────────────────────────────────────┘'
    )

    # Prerequisites
    add_heading(doc, 'Prerequisites')
    add_table(doc,
        ['Requirement', 'Details'],
        [
            ['AWS Organizations', 'Configured with OU structure'],
            ['Delegated Backup Admin', 'Account registered for AWS Backup'],
            ['AWS CLI v2', 'Admin credentials for delegated admin'],
            ['Management account', 'Access for one-time setup (Step 1)'],
            ['Amazon DevOps Agent', 'GA access'],
            ['Slack workspace', 'Permission to install apps'],
        ]
    )

    # Step 1
    add_heading(doc, 'Step 1: One-Time Setup (Management Account)')
    doc.add_paragraph('Register your delegated admin for CloudFormation StackSets:')
    add_code_block(doc,
        'aws organizations register-delegated-administrator \\\n'
        '  --account-id YOUR_DELEGATED_ADMIN_ID \\\n'
        '  --service-principal member.org.stacksets.cloudformation.amazonaws.com'
    )
    p = doc.add_paragraph()
    p.add_run('✅ Verify: ').bold = True
    add_code_block(doc,
        'aws organizations list-delegated-administrators \\\n'
        '  --service-principal member.org.stacksets.cloudformation.amazonaws.com'
    )
    doc.add_paragraph('You will not need management account access again.').italic = True

    # Step 2
    add_heading(doc, 'Step 2: Deploy the Solution (Delegated Admin Account)')

    add_heading(doc, 'Step 2.1: Clone and Run', level=2)
    add_code_block(doc,
        'git clone https://github.com/kukkunuruanil/aws-backup-devops-agent-v6.git\n'
        'cd aws-backup-devops-agent\n'
        './deploy.sh'
    )

    add_heading(doc, 'Step 2.2: Prompts (3 inputs only)', level=2)
    add_table(doc,
        ['Prompt', 'Example', 'Where to find it'],
        [
            ['Organization ID', 'o-abc1234567', 'aws organizations describe-organization'],
            ['Target OU ID', 'ou-xxxx-xxxxxxxx', 'AWS Organizations console'],
            ['Region', 'us-west-2', 'Your primary backup region'],
        ]
    )

    add_heading(doc, 'Step 2.3: Automated Steps', level=2)
    doc.add_paragraph('The script performs 5 steps automatically:')
    add_table(doc,
        ['Step', 'Action', 'Creates'],
        [
            ['1/5', 'Create Agent Space & Webhook', 'DevOps Agent space + HMAC webhook'],
            ['2/5', 'Deploy main stack', 'Lambda, EventBridge, IAM, Secrets Manager'],
            ['3/5', 'Associate account', 'Links AWS account to Agent Space'],
            ['4/5', 'Deploy StackSet', 'Event forwarder + investigation role in all members'],
            ['5/5', 'Verify deployment', 'Per-account status table + Lambda test'],
        ]
    )

    add_heading(doc, 'Step 2.4: Verify StackSet Deployment', level=2)
    doc.add_paragraph('The script automatically shows verification:')
    add_code_block(doc,
        '┌─────────────────┬──────────┬───────────────────────┐\n'
        '│ Account         │ Region   │ Status                │\n'
        '├─────────────────┼──────────┼───────────────────────┤\n'
        '│ 444455556666    │ us-west-2│ ✓ SUCCEEDED           │\n'
        '│ 777788889999    │ us-west-2│ ✓ SUCCEEDED           │\n'
        '└─────────────────┴──────────┴───────────────────────┘\n\n'
        '→ Testing Lambda (sending synthetic failure event)...\n'
        '✓ Lambda invocation: HTTP 200 | Response: {"statusCode": 200}\n\n'
        '╔════════════════════════════════════════════════════════════╗\n'
        '║              ✓ DEPLOYMENT COMPLETE                        ║\n'
        '╚════════════════════════════════════════════════════════════╝'
    )
    doc.add_paragraph('Manual verification (if needed):')
    add_code_block(doc,
        'aws cloudformation list-stack-instances \\\n'
        '  --stack-set-name BackupEventForwarder \\\n'
        '  --call-as DELEGATED_ADMIN --region us-west-2 \\\n'
        '  --query \'Summaries[].{Account:Account,Status:StackInstanceStatus.DetailedStatus}\''
    )

    # Step 3
    add_heading(doc, 'Step 3: Connect Slack')

    add_heading(doc, 'Step 3.1: Register Slack', level=2)
    doc.add_paragraph('DevOps Agent console → Capability Providers → Slack → Register')
    doc.add_paragraph('Complete OAuth authorization.')

    add_heading(doc, 'Step 3.2: Associate Slack Channel', level=2)
    add_code_block(doc,
        'aws devops-agent list-services --region us-west-2\n\n'
        'aws devops-agent associate-service \\\n'
        '  --agent-space-id YOUR_AGENT_SPACE_ID \\\n'
        '  --service-id YOUR_SLACK_SERVICE_UUID \\\n'
        '  --configuration \'{"slack":{"workspaceId":"YOUR_TEAM_ID",\\\n'
        '    "workspaceName":"YOUR_WORKSPACE",\\\n'
        '    "transmissionTarget":{"opsOncallTarget":{"channelId":"YOUR_CHANNEL_ID"}}}}\' \\\n'
        '  --region us-west-2'
    )

    add_heading(doc, 'Step 3.3: Add Bot to Channel', level=2)
    doc.add_paragraph('/invite @AWS DevOps Agent')

    # Step 4
    add_heading(doc, 'Step 4: Test End-to-End')
    doc.add_paragraph('Synthetic test (instant):')
    add_code_block(doc,
        'aws lambda invoke --function-name BackupFailureBridge \\\n'
        '  --payload $(echo \'{"source":"aws.backup","detail-type":"Copy Job State Change",\\\n'
        '    "detail":{"state":"FAILED","copyJobId":"TEST-001",\\\n'
        '    "statusMessage":"KMS key not accessible",\\\n'
        '    "accountId":"111122223333",\\\n'
        '    "resourceArn":"arn:aws:ec2:us-west-2:111122223333:volume/vol-0abc",\\\n'
        '    "backupVaultName":"Default"}}\' | base64) \\\n'
        '  /tmp/test.json && cat /tmp/test.json'
    )
    doc.add_paragraph('Expected: {"statusCode": 200} + investigation in Slack within 3–5 minutes.')

    # What gets deployed
    add_heading(doc, 'What Gets Deployed')

    add_heading(doc, 'Delegated Admin Account', level=2)
    add_table(doc,
        ['Resource', 'Type', 'Purpose'],
        [
            ['BackupInvestigations', 'Agent Space', 'Receives webhook events'],
            ['HMAC Webhook', 'Webhook', 'Authenticated event intake'],
            ['BackupFailureBridge', 'Lambda', 'Transforms events → webhook POST'],
            ['BackupFailures-TriggerInvestigation', 'EventBridge Rule', 'Catches FAILED events'],
            ['EventBus Policy', 'EventBridge', 'Allows org accounts to send events'],
            ['DevOpsAgentBackupRole', 'IAM Role', 'Agent investigation access'],
            ['BackupFailureBridgeRole', 'IAM Role', 'Lambda execution'],
            ['devops-agent/backup-webhook', 'Secret', 'Webhook URL + HMAC key'],
        ]
    )

    add_heading(doc, 'Member Accounts (via StackSet)', level=2)
    add_table(doc,
        ['Resource', 'Type', 'Purpose'],
        [
            ['ForwardBackupFailures-ToCentral', 'EventBridge Rule', 'Forwards FAILED events'],
            ['BackupEventForwardingRole', 'IAM Role', 'Cross-account event delivery'],
            ['DevOpsAgentInvestigationRole', 'IAM Role', 'Agent cross-account analysis'],
        ]
    )

    # Troubleshooting
    add_heading(doc, 'Troubleshooting')
    add_table(doc,
        ['Issue', 'Check'],
        [
            ['StackSet creation fails', 'Registered as CF StackSets delegated admin? (Step 1)'],
            ['No events from members', 'list-stack-instances --call-as DELEGATED_ADMIN'],
            ['Lambda error', 'aws logs tail /aws/lambda/BackupFailureBridge --since 10m'],
            ['No Slack notification', 'Bot in channel? Check Agent console'],
            ['Webhook 403', 'Secret rotated? Update Secrets Manager'],
            ['Agent can\'t investigate', 'DevOpsAgentInvestigationRole trusts delegated admin?'],
        ]
    )

    # Cleanup
    add_heading(doc, 'Cleanup')
    add_code_block(doc,
        'REGION=us-west-2\n\n'
        '# 1. Remove StackSet instances\n'
        'aws cloudformation delete-stack-instances \\\n'
        '  --stack-set-name BackupEventForwarder \\\n'
        '  --deployment-targets OrganizationalUnitIds=YOUR_OU_ID \\\n'
        '  --regions $REGION --no-retain-stacks \\\n'
        '  --call-as DELEGATED_ADMIN --region $REGION\n\n'
        '# 2. Delete StackSet\n'
        'aws cloudformation delete-stack-set \\\n'
        '  --stack-set-name BackupEventForwarder \\\n'
        '  --call-as DELEGATED_ADMIN --region $REGION\n\n'
        '# 3. Delete main stack\n'
        'aws cloudformation delete-stack --stack-name BackupDevOpsAgent --region $REGION\n\n'
        '# 4. (Optional) Delete Agent Space\n'
        'aws devops-agent delete-agent-space \\\n'
        '  --agent-space-id YOUR_AGENT_SPACE_ID --region $REGION'
    )

    path = os.path.join(OUTPUT_DIR, 'Implementation_Guide_AWS_Backup_DevOps_Agent.docx')
    doc.save(path)
    print(f'✓ Implementation guide: {path}')
    return path


if __name__ == '__main__':
    generate_blog_post()
    generate_implementation_guide()
    print('\nDone! Both DOCX files generated.')
